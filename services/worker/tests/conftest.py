from pathlib import Path

import pytest
from docx import Document

from app.config import Settings


@pytest.fixture()
def settings() -> Settings:
    return Settings(
        redis_url="redis://localhost:6379/0",
        mongo_uri="mongodb://localhost:27017",
        mongo_db="docuparse_test",
        upload_dir="/data/uploads",
        groq_api_key="test-key",
        groq_model="llama-3.1-8b-instant",
    )


@pytest.fixture()
def pdf_file(tmp_path: Path) -> Path:
    # Minimal valid PDF with one text-bearing page
    minimal_pdf = (
        b"%PDF-1.4\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792]\n"
        b"   /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        b"4 0 obj\n<< /Length 44 >>\nstream\n"
        b"BT /F1 12 Tf 72 720 Td (Jane Doe Resume) Tj ET\n"
        b"endstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000360 00000 n \n"
        b"trailer\n<< /Size 6 /Root 1 0 R >>\nstartxref\n441\n%%EOF\n"
    )
    path = tmp_path / "resume.pdf"
    path.write_bytes(minimal_pdf)
    return path


@pytest.fixture()
def docx_file(tmp_path: Path) -> Path:
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Skills: Python, FastAPI")
    path = tmp_path / "resume.docx"
    doc.save(str(path))
    return path
